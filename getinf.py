#!/usr/bin/env python3

import os
from docx import Document
import fitz  # PyMuPDF
import PyPDF2
from PIL import Image
import pytesseract


def extract_text_from_pdf(pdf_path):
    text = ''
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in range(len(reader.pages)):
                text += reader.pages[page].extract_text()
    except Exception as e:
        print(f"Error extracting text from PDF '{pdf_path}': {e}")
    return text


def extract_text_from_docx(docx_path):
    text = ''
    try:
        document = Document(docx_path) 
        col_keys = []       # Column name
        col_values = []     # Column value
        index_num = 0
        # Extract form information, add a deduplication mechanism
        fore_str = ""
        cell_text = ""
        for table in document.tables:
            for row_index,row in enumerate(table.rows):
                for col_index,cell in enumerate(row.cells):
                    if fore_str != cell.text:
                        if index_num % 2==0:
                            col_keys.append(cell.text)
                        else:
                            col_values.append(cell.text)
                        fore_str = cell.text
                        index_num +=1
                        cell_text += cell.text + '\n'
        # Extract text information
        paragraphs_text = ""
        for paragraph in document.paragraphs:
            paragraphs_text += paragraph.text + "\n"

        text = cell_text + paragraphs_text
    except Exception as e:
        print(f"Error extracting text from DOCX '{docx_path}': {e}")
    return text


def extract_text_from_png(docx_path):
    # Open image
    img = Image.open(docx_path)

    # Character recognition
    text = pytesseract.image_to_string(img, lang = 'chi_sim+eng_sim')

    text = ''.join([i.strip(' ') for i in text])
    # Print results
    print("识别结果：", text)

    return text


def save_text_to_file(text, file_path):
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(text)
        print(f"Text saved to '{file_path}'")
    except Exception as e:
        print(f"Error saving text to file '{file_path}': {e}")


def batch_extract_text_from_files(directory):
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        elif filename.endswith('.png'):
            text = extract_text_from_png(file_path)
        else:
            continue  # Ignore files with unsupported formats
        
        # Save extracted text to a file with the same name
        text_file_path = os.path.splitext(file_path)[0] + '_text.txt'
        save_text_to_file(text, text_file_path)


def main():
    batch_extract_text_from_files('./file_and_text/')  # Extract text from files in the current directory

if __name__ == "__main__":
    main()
